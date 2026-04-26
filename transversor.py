"""
Transversor PDF - Document Translation Tool
-------------------------------------------
This application allows users to translate PDF documents into multiple languages
while maintaining a professional output format in both PDF and DOCX.

Main features:
- PDF text extraction using PyMuPDF (fitz).
- Multi-language translation via Google Translator.
- Selection of specific pages for processing.
- Export to Word (.docx) and PDF formats.

Author: AlvaroEC
License: MIT
"""

import os
import sys
import ctypes
import threading
import textwrap
import re
import customtkinter as ctk
from tkinter import messagebox, filedialog
from tkinterdnd2 import DND_FILES, TkinterDnD
from PIL import Image, ImageTk
from PyPDF2 import PdfReader
from deep_translator import GoogleTranslator
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from docx import Document
import fitz  
from PIL import Image

ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

def resource_path(relative_path):
    """
    Adjusts file paths for bundled executable resources.

    Args:
        relative_path (str): The relative path to the resource file.

    Returns:
        str: The absolute path to the resource, compatible with PyInstaller.
    """
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

try:
    ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID('transversor.pdf')
except:
    pass

class App(TkinterDnD.Tk):
    """
    Main Application class for Transversor PDF.
    
    Inherits from TkinterDnD.Tk to support Drag & Drop functionality.
    Manages the GUI, file queue, and the translation orchestration logic.
    """
    def __init__(self):
        """
        Initializes the application window, state variables, and UI components.
        """
        super().__init__()
        
        self.title("Transversor PDF")
        self.geometry("1120x700")
        self.minsize(980, 620)

        self.archivo_icono = resource_path("icono.ico")
        if os.path.exists(self.archivo_icono):
            self.iconbitmap(self.archivo_icono)
            try:
                img = Image.open(self.archivo_icono)
                photo = ImageTk.PhotoImage(img)
                self.iconphoto(False, photo)
                self._icon_photo = photo
            except Exception:
                pass
        
        # --- STATE ---
        self.cola_archivos = {} 
        self.ruta_destino = os.path.join(os.path.expanduser('~'), 'Desktop')
        self.archivo_editando = None
        self.progress_window = None
        self.progress_bar_popup = None
        self.progress_label_popup = None

        # --- UI LAYOUT ---
        self._setup_ui()

    def _setup_ui(self):
        """
        Internal method to build the sidebar and the main content panels.
        """
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        # 1. SIDEBAR
        self.sidebar = ctk.CTkFrame(self, width=280, corner_radius=0)
        self.sidebar.grid(row=0, column=0, sticky="nsew")
        self.sidebar.grid_propagate(False)

        ctk.CTkLabel(self.sidebar, text="TRANSVERSOR PDF", font=("Arial", 22, "bold")).pack(pady=(16, 2))
        ctk.CTkLabel(self.sidebar, text="Translation Suite", text_color="lightgray", font=("Arial", 11)).pack(pady=(0, 14))
        
        # Section 1: Loading
        ctk.CTkLabel(self.sidebar, text="1) LOAD", font=("Arial", 12, "bold")).pack(pady=(2, 5))
        ctk.CTkButton(self.sidebar, text="📁 Add PDFs", command=self.examinar_archivos, width=230).pack(pady=4, padx=20)
        ctk.CTkButton(self.sidebar, text="📂 Add Folder", command=self.examinar_carpeta, width=230).pack(pady=4, padx=20)
        
        # Section 2: Destination
        ctk.CTkLabel(self.sidebar, text="2) DESTINATION", font=("Arial", 12, "bold")).pack(pady=(14, 5))
        self.btn_dest = ctk.CTkButton(self.sidebar, text="📍 Choose Output Folder", command=self.elegir_destino, fg_color="#8e44ad", width=230)
        self.btn_dest.pack(pady=4, padx=20)
        self.lbl_destino = ctk.CTkLabel(self.sidebar, text=f"Path:\n{self.ruta_destino}", text_color="lightgray", justify="left", wraplength=240, anchor="w")
        self.lbl_destino.pack(padx=20, pady=(5, 8), fill="x")
        self.btn_abrir_destino = ctk.CTkButton(self.sidebar, text="👁️ Open Folder", command=self.abrir_destino, fg_color="#34495e", width=230)
        self.btn_abrir_destino.pack(pady=(0, 12), padx=20)

        # Section 3: Translation Options
        self.idiomas = {"Español": "es", "Inglés": "en", "Francés": "fr", "Italiano": "it", "Portugués": "pt", "Alemán": "de", "Chino": "zh-CN"}
        ctk.CTkLabel(self.sidebar, text="3) TRANSLATION", font=("Arial", 12, "bold")).pack(pady=(2, 5))
        self.combo_idioma = ctk.CTkComboBox(self.sidebar, values=list(self.idiomas.keys()))
        self.combo_idioma.set("Español")
        self.combo_idioma.pack(pady=(4, 12), padx=20)

        self.formato_var = ctk.StringVar(value="PDF")
        ctk.CTkRadioButton(self.sidebar, text="Export to PDF", variable=self.formato_var, value="PDF").pack()
        ctk.CTkRadioButton(self.sidebar, text="Export to Word", variable=self.formato_var, value="DOCX").pack(pady=(4, 8))

        # Section 4: Structure
        ctk.CTkLabel(self.sidebar, text="STRUCTURE:", font=("Arial", 12, "bold")).pack(pady=(10, 5))
        self.estructura_var = ctk.StringVar(value="ESTRICTO")
        ctk.CTkRadioButton(self.sidebar, text="Continuous block", variable=self.estructura_var, value="CONTINUO").pack()
        ctk.CTkRadioButton(self.sidebar, text="Page by page", variable=self.estructura_var, value="ESTRICTO").pack(pady=(5, 0))

        # Actions
        self.btn_traducir = ctk.CTkButton(self.sidebar, text="🚀 Translate All", fg_color="#2ecc71", command=self.iniciar_proceso, height=48, width=230, font=("Arial", 14, "bold"))
        self.btn_traducir.pack(pady=(18, 10), padx=20)
        ctk.CTkButton(self.sidebar, text="🧹 Clear List", fg_color="#e74c3c", command=self.limpiar_lista, width=230).pack()

        self.status = ctk.CTkLabel(self.sidebar, text="Ready • Double-click to edit", text_color="cyan", wraplength=240)
        self.status.pack(pady=16, padx=20)

        # RIGHT PANEL
        self.right_container = ctk.CTkFrame(self, fg_color="transparent")
        self.right_container.grid(row=0, column=1, sticky="nsew", padx=10, pady=10)
        self.right_container.grid_columnconfigure(0, weight=1)
        self.right_container.grid_rowconfigure(0, weight=1)

        self.file_viewer = ctk.CTkScrollableFrame(self.right_container, label_text="DOCUMENTS • Drag & Drop PDFs here")
        self.file_viewer.grid(row=0, column=0, sticky="nsew")
        self.file_viewer.grid_columnconfigure((0, 1, 2), weight=1)

        self.page_selector = ctk.CTkFrame(self.right_container)

        self.drop_target_register(DND_FILES)
        self.dnd_bind('<<Drop>>', self.al_soltar)

    def mostrar_selector(self, path):
        """
        Switches the UI to the page selection panel for a specific PDF.

        Args:
            path (str): The file path of the PDF to edit.
        """
        self.archivo_editando = path
        self.file_viewer.grid_forget()
        self.page_selector.grid(row=0, column=0, sticky="nsew")
        self.dibujar_selector_paginas(path)

    def volver_a_lista(self):
        """
        Switches the UI back to the main file queue viewer.
        """
        self.page_selector.grid_forget()
        self.file_viewer.grid(row=0, column=0, sticky="nsew")
        self.archivo_editando = None
        self.actualizar_visor_archivos()

    def agregar_a_cola(self, path):
        """
        Adds a new PDF to the internal queue and initializes the page selection.

        Args:
            path (str): The absolute path to the PDF file.
        """
        if path not in self.cola_archivos:
            doc = fitz.open(path)
            self.cola_archivos[path] = list(range(len(doc)))
            doc.close()

    def al_soltar(self, event):
        """
        Handles Drag & Drop events for files and folders.

        Args:
            event: The TkinterDnD event containing the dropped file paths.
        """
        paths = [p[0] if p[0] else p[1] for p in re.findall(r'\{(.*?)\}|(\S+)', event.data)]
        for p in paths:
            p = p.strip('{}').replace('"', '')
            if os.path.isdir(p):
                for f in os.listdir(p):
                    if f.lower().endswith('.pdf'): self.agregar_a_cola(os.path.join(p, f))
            elif p.lower().endswith('.pdf'): self.agregar_a_cola(p)
        self.actualizar_visor_archivos()

    def actualizar_visor_archivos(self):
        """
        Refreshes the file viewer panel with thumbnails and document info.
        """
        for widget in self.file_viewer.winfo_children(): widget.destroy()

        if not self.cola_archivos:
            estado_vacio = ctk.CTkFrame(self.file_viewer, fg_color="#2c3e50")
            estado_vacio.grid(row=0, column=0, columnspan=3, sticky="ew", padx=24, pady=24)
            ctk.CTkLabel(estado_vacio, text="📄 No PDFs loaded", font=("Arial", 16, "bold")).pack(pady=(18, 8))
            ctk.CTkLabel(estado_vacio, text="Drag files here or use sidebar buttons.", text_color="lightgray").pack(pady=(0, 18))
            return

        for i, (path, seleccion) in enumerate(self.cola_archivos.items()):
            frame = ctk.CTkFrame(self.file_viewer, fg_color="#2c3e50")
            frame.grid(row=i//3, column=i%3, padx=10, pady=10, sticky="nsew")
            
            doc = fitz.open(path)
            pix = doc[0].get_pixmap(matrix=fitz.Matrix(0.1, 0.1))
            img = ctk.CTkImage(Image.frombytes("RGB", [pix.width, pix.height], pix.samples), size=(100, 130))
            
            lbl_img = ctk.CTkLabel(frame, image=img, text="")
            lbl_img.pack(pady=5)
            ctk.CTkLabel(frame, text=os.path.basename(path), font=("Arial", 10, "bold"), wraplength=140).pack(padx=6)
            ctk.CTkLabel(frame, text=f"Pages: {len(seleccion)}/{len(doc)}", font=("Arial", 9), text_color="gray").pack()
            doc.close()

            frame.bind("<Double-Button-1>", lambda e, p=path: self.mostrar_selector(p))
            lbl_img.bind("<Double-Button-1>", lambda e, p=path: self.mostrar_selector(p))

    def dibujar_selector_paginas(self, path):
        """
        Renders the page selector editor for a specific document.

        Args:
            path (str): File path of the document being edited.
        """
        for widget in self.page_selector.winfo_children(): widget.destroy()
        
        lbl_info = ctk.CTkLabel(self.page_selector, text=f"Edit Selection • {os.path.basename(path)}", font=("Arial", 14, "bold"))
        lbl_info.pack(pady=(10, 4))

        scroll_p = ctk.CTkScrollableFrame(self.page_selector)
        scroll_p.pack(fill="both", expand=True, padx=10, pady=5)
        scroll_p.grid_columnconfigure((0,1,2,3), weight=1)

        doc = fitz.open(path)
        checks = {}
        for i in range(len(doc)):
            f_p = ctk.CTkFrame(scroll_p, fg_color="#333333")
            f_p.grid(row=i//4, column=i%4, padx=5, pady=5)
            
            pix = doc[i].get_pixmap(matrix=fitz.Matrix(0.08, 0.08))
            img = ctk.CTkImage(Image.frombytes("RGB", [pix.width, pix.height], pix.samples), size=(80, 110))
            
            ctk.CTkLabel(f_p, image=img, text="").pack()
            var = ctk.BooleanVar(value=i in self.cola_archivos[path])
            checks[i] = var
            ctk.CTkCheckBox(f_p, text=f"Pg {i+1}", variable=var).pack(pady=2)

        def confirmar():
            self.cola_archivos[path] = sorted({pg for pg, v in checks.items() if v.get()})
            doc.close()
            self.volver_a_lista()

        ctk.CTkButton(self.page_selector, text="✅ Save Changes", fg_color="#2ecc71", command=confirmar, height=42).pack(pady=10)

    def iniciar_proceso(self):
        """
        Captures the UI configuration and launches the translation thread.
        """
        if not self.cola_archivos:
            return
        config = {
            "idioma_destino": self.idiomas[self.combo_idioma.get()],
            "formato": self.formato_var.get(),
            "separar_paginas": self.estructura_var.get() == "ESTRICTO"
        }
        self._mostrar_ventana_progreso_ui()
        self.btn_traducir.configure(state="disabled")
        threading.Thread(target=self.hilo_traduccion, args=(config,), daemon=True).start()

    def _set_status_ui(self, texto, color="gray"):
        """Updates the status label in the UI thread."""
        self.after(0, lambda: self.status.configure(text=texto, text_color=color))

    def _mostrar_ventana_progreso_ui(self):
        """Creates a modal popup window to show the translation progress bar."""
        if self.progress_window and self.progress_window.winfo_exists():
            return
        win = ctk.CTkToplevel(self)
        win.title("Progress")
        win.geometry("420x150")
        win.resizable(False, False)
        win.attributes("-topmost", True)
        
        self.update_idletasks()
        x = self.winfo_rootx() + (self.winfo_width() // 2) - 210
        y = self.winfo_rooty() + (self.winfo_height() // 2) - 75
        win.geometry(f"420x150+{x}+{y}")

        ctk.CTkLabel(win, text="Translating Documents", font=("Arial", 16, "bold")).pack(pady=(18, 6))
        self.progress_label_popup = ctk.CTkLabel(win, text="Processing... 0%", text_color="lightgray")
        self.progress_label_popup.pack(pady=(0, 8))
        self.progress_bar_popup = ctk.CTkProgressBar(win, width=340)
        self.progress_bar_popup.pack(pady=(0, 14))
        self.progress_bar_popup.set(0)
        win.grab_set()
        self.progress_window = win

    def _cerrar_ventana_progreso_ui(self):
        """Destroys the progress popup window."""
        if self.progress_window and self.progress_window.winfo_exists():
            self.progress_window.destroy()
        self.progress_window = None

    def _actualizar_progreso_popup_ui(self, valor):
        """Updates the progress bar value and label in the popup."""
        if self.progress_bar_popup and self.progress_bar_popup.winfo_exists():
            self.progress_bar_popup.set(valor)
        if self.progress_label_popup and self.progress_label_popup.winfo_exists():
            self.progress_label_popup.configure(text=f"Progress... {int(valor * 100)}%")

    def _set_progress_ui(self, valor):
        """Internal helper to safely update progress from a thread."""
        v = max(0.0, min(1.0, valor))
        self.after(0, lambda: self._actualizar_progreso_popup_ui(v))

    def hilo_traduccion(self, config):
        """
        Core translation engine executed in a background thread.
        Handles text extraction, chunking, translation, and file generation.

        Args:
            config (dict): Configuration containing language, format, and structure settings.
        """
        try:
            self._set_progress_ui(0)
            translator = GoogleTranslator(source='auto', target=config["idioma_destino"])
            archivos = list(self.cola_archivos.items())
            total_bloques = 0
            estructura_archivos = []

            # Preparation Phase: Calculate total work
            for p_in, paginas in archivos:
                if not paginas: continue
                reader = PdfReader(p_in)
                doc_fitz = fitz.open(p_in)
                paginas_origen = []
                for idx in sorted(set(paginas)):
                    parrafos = self.extraer_parrafos_pagina(doc_fitz, idx)
                    if not parrafos:
                        t = reader.pages[idx].extract_text()
                        parrafos = self.extraer_parrafos(t) if t else []
                    paginas_origen.append(parrafos)
                    for p in parrafos:
                        total_bloques += max(1, len(self.construir_bloques_traduccion(p)))
                estructura_archivos.append((p_in, paginas_origen))
                doc_fitz.close()

            if total_bloques <= 0: total_bloques = 1
            bloques_completados = 0

            # Processing Phase: Translate and Save
            for i, (p_in, paginas_origen) in enumerate(estructura_archivos):
                self._set_status_ui(f"Translating {i+1}/{len(estructura_archivos)}...", "orange")
                paginas_traducidas = []
                
                def avanzar_progreso():
                    nonlocal bloques_completados
                    bloques_completados += 1
                    self._set_progress_ui(bloques_completados / total_bloques)

                for parrafos in paginas_origen:
                    parrafos_traducidos = [self.traducir_texto(translator, p, progreso_callback=avanzar_progreso) for p in parrafos]
                    paginas_traducidas.append(parrafos_traducidos)

                ext = ".pdf" if config["formato"] == "PDF" else ".docx"
                p_out = os.path.join(self.ruta_destino, f"Translated_{os.path.basename(p_in).replace('.pdf', ext)}")
                
                if config["formato"] == "PDF":
                    self.guardar_pdf_estructurado(paginas_traducidas, p_out, separar_paginas=config["separar_paginas"])
                else:
                    self.guardar_docx(paginas_traducidas, p_out, separar_paginas=config["separar_paginas"])

            self._set_progress_ui(1)
            self._set_status_ui("COMPLETED!", "green")
            self.after(0, lambda: messagebox.showinfo("Success", "Files processed correctly."))
        except Exception as e:
            self._set_status_ui("Translation Error", "red")
            self.after(0, lambda: messagebox.showerror("Error", f"Failed:\n{e}"))
        finally:
            self.after(0, self._cerrar_ventana_progreso_ui)
            self.after(0, lambda: self.btn_traducir.configure(state="normal"))

    def guardar_pdf_estructurado(self, paginas_texto, path, separar_paginas=True):
        """
        Generates a professional PDF with basic structure and text wrapping.

        Args:
            paginas_texto (list): Nested list of translated paragraphs by page.
            path (str): Destination file path.
            separar_paginas (bool): Whether to respect original page breaks.
        """
        c = canvas.Canvas(path, pagesize=A4)
        text_obj = c.beginText(50, 800)
        text_obj.setFont("Helvetica", 10)
        pagina_vacia = True

        for i, parrafos_pagina in enumerate(paginas_texto):
            if parrafos_pagina:
                for parrafo in parrafos_pagina:
                    texto = parrafo.strip()
                    if not texto: continue
                    lineas = textwrap.wrap(texto, width=95)
                    for linea in lineas:
                        if text_obj.getY() < 50:
                            c.drawText(text_obj); c.showPage()
                            text_obj = c.beginText(50, 800); text_obj.setFont("Helvetica", 10)
                            pagina_vacia = True
                        text_obj.textLine(linea)
                        pagina_vacia = False
                    text_obj.textLine("") 

            if separar_paginas and i < len(paginas_texto) - 1:
                if not pagina_vacia:
                    c.drawText(text_obj); c.showPage()
                    text_obj = c.beginText(50, 800); text_obj.setFont("Helvetica", 10)
                    pagina_vacia = True

        c.drawText(text_obj); c.save()

    def construir_bloques_traduccion(self, texto, tam_bloque=1800):
        """
        Splits a paragraph into smaller chunks to respect API character limits.

        Args:
            texto (str): The full paragraph string.
            tam_bloque (int): Maximum character count per chunk.

        Returns:
            list: List of text chunks.
        """
        contenido = texto.strip()
        if not contenido: return []
        oraciones = self.dividir_en_oraciones(contenido)
        bloques, actual = [], ""
        for oracion in oraciones:
            candidato = (actual + " " + oracion).strip() if actual else oracion
            if len(candidato) <= tam_bloque: actual = candidato
            else:
                if actual: bloques.append(actual)
                actual = oracion
        if actual: bloques.append(actual)
        return bloques

    def traducir_texto(self, translator, texto, tam_bloque=1800, progreso_callback=None):
        """
        Orchestrates the translation of a single paragraph through chunks.

        Args:
            translator: The GoogleTranslator instance.
            texto (str): Paragraph to translate.
            tam_bloque (int): Max chunk size.
            progreso_callback (callable): Function to call after each successful chunk translation.

        Returns:
            str: The fully translated paragraph.
        """
        if not texto.strip(): return ""
        prefijo = ""
        contenido = texto.strip()
        if contenido.startswith(("- ", "• ", "* ")):
            prefijo = contenido[:2]; contenido = contenido[2:].strip()
        
        bloques = self.construir_bloques_traduccion(contenido, tam_bloque=tam_bloque)
        traducido = []
        for bloque in bloques:
            try: traducido.append(translator.translate(bloque))
            except: traducido.append(bloque)
            if progreso_callback: progreso_callback()

        res = self.limpiar_texto_traducido(" ".join(traducido))
        return f"{prefijo}{res}" if prefijo else res

    def dividir_en_oraciones(self, texto):
        """Uses regex to split text into sentences based on punctuation."""
        texto_norm = re.sub(r"\s+", " ", texto).strip()
        return [p.strip() for p in re.split(r'(?<=[.!?;:])\s+(?=[A-ZÁÉÍÓÚÜÑ0-9"“(])', texto_norm) if p.strip()]

    def limpiar_texto_traducido(self, texto):
        """Sanitizes the translated string removing redundant spaces and fixing punctuation."""
        t = re.sub(r"\s+", " ", texto.strip())
        t = re.sub(r"\s+([,.;:!?])", r"\1", t)
        return t

    def extraer_parrafos_pagina(self, doc_fitz, page_index):
        """
        Intelligent text extraction using visual blocks to identify paragraphs.

        Args:
            doc_fitz: PyMuPDF document instance.
            page_index (int): Page number to extract.

        Returns:
            list: List of identified paragraphs.
        """
        page = doc_fitz[page_index]
        blocks = sorted(page.get_text("dict").get("blocks", []), key=lambda b: (round(b.get("bbox", [0, 0])[1], 1), round(b.get("bbox", [0, 0])[0], 1)))
        parrafos = []
        for bloque in blocks:
            if bloque.get("type") != 0: continue
            actual = ""
            for linea in bloque.get("lines", []):
                t_linea = " ".join([s.get("text", "").strip() for s in linea.get("spans", []) if s.get("text", "").strip()]).strip()
                if not t_linea: continue
                if not actual: actual = t_linea
                else: actual = (actual[:-1] + t_linea) if actual.endswith("-") else (actual + " " + t_linea)
            if actual.strip(): parrafos.append(actual.strip())
        return parrafos

    def guardar_docx(self, paginas_texto, path, separar_paginas=True):
        """
        Generates a Word document with bullet points and numbered list detection.

        Args:
            paginas_texto (list): Translated content.
            path (str): Target path.
            separar_paginas (bool): Respect page breaks.
        """
        doc = Document()
        for i, parrafos_pagina in enumerate(paginas_texto):
            if separar_paginas and i > 0: doc.add_page_break()
            for parrafo in parrafos_pagina:
                t = parrafo.strip()
                if not t: continue
                if t.startswith(("- ", "• ")): doc.add_paragraph(t[2:].strip(), style="List Bullet")
                elif len(t) > 2 and t[0].isdigit() and t[1:3] in (". ", ") "): doc.add_paragraph(t[3:].strip(), style="List Number")
                else: doc.add_paragraph(t)
        doc.save(path)

    def examinar_archivos(self):
        """Opens a file dialog to select multiple PDF files."""
        files = filedialog.askopenfilenames(filetypes=[("PDF", "*.pdf")])
        if files:
            for f in files: self.agregar_a_cola(f)
            self.actualizar_visor_archivos()

    def examinar_carpeta(self):
        """Opens a directory dialog to scan for PDF files."""
        folder = filedialog.askdirectory()
        if folder:
            for f in os.listdir(folder):
                if f.lower().endswith('.pdf'): self.agregar_a_cola(os.path.join(folder, f))
            self.actualizar_visor_archivos()

    def elegir_destino(self):
        """Updates the output directory path."""
        dest = filedialog.askdirectory()
        if dest:
            self.ruta_destino = dest
            self.lbl_destino.configure(text=f"Path:\n{self.ruta_destino}")

    def abrir_destino(self):
        """Opens the selected output folder in the OS file explorer."""
        if os.path.isdir(self.ruta_destino): os.startfile(self.ruta_destino)
        else: messagebox.showerror("Error", "Folder does not exist.")

    def limpiar_lista(self):
        """Clears all documents from the current queue."""
        self.cola_archivos = {}; self.actualizar_visor_archivos()

if __name__ == "__main__":
    app = App()
    app.mainloop()
