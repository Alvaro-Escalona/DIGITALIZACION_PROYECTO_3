import os
import sys
import ctypes
import threading
import textwrap
import re
import customtkinter as ctk
from tkinter import messagebox, filedialog
from tkinterdnd2 import DND_FILES, TkinterDnD
from PIL import Image, ImageTk
from PyPDF2 import PdfReader
from deep_translator import GoogleTranslator
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from docx import Document
import fitz  
from PIL import Image

ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

try:
    ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID('transversor.pdf')
except:
    pass

class App(TkinterDnD.Tk):
    def __init__(self):
        super().__init__()
        
        self.title("Transversor PDF")
        self.geometry("1120x700")
        self.minsize(980, 620)

        self.archivo_icono = resource_path("icono.ico")
        if os.path.exists(self.archivo_icono):
            self.iconbitmap(self.archivo_icono)
            try:
                img = Image.open(self.archivo_icono)
                photo = ImageTk.PhotoImage(img)
                self.iconphoto(False, photo)
                self._icon_photo = photo
            except Exception:
                pass
        
        # --- ESTADO ---
        self.cola_archivos = {} 
        self.ruta_destino = os.path.join(os.path.expanduser('~'), 'Desktop')
        self.archivo_editando = None
        self.progress_window = None
        self.progress_bar_popup = None
        self.progress_label_popup = None

        # --- GRID PRINCIPAL ---
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        # 1. PANEL IZQUIERDO
        self.sidebar = ctk.CTkFrame(self, width=280, corner_radius=0)
        self.sidebar.grid(row=0, column=0, sticky="nsew")
        self.sidebar.grid_propagate(False)

        ctk.CTkLabel(self.sidebar, text="TRANSVERSOR PDF", font=("Arial", 22, "bold")).pack(pady=(16, 2))
        ctk.CTkLabel(
            self.sidebar,
            text="Transversor PDF",
            text_color="lightgray",
            font=("Arial", 11)
        ).pack(pady=(0, 14))
        
        ctk.CTkLabel(self.sidebar, text="1) CARGA", font=("Arial", 12, "bold")).pack(pady=(2, 5))
        ctk.CTkButton(self.sidebar, text="📁 Añadir PDFs", command=self.examinar_archivos, width=230).pack(pady=4, padx=20)
        ctk.CTkButton(self.sidebar, text="📂 Añadir carpeta", command=self.examinar_carpeta, width=230).pack(pady=4, padx=20)
        
        ctk.CTkLabel(self.sidebar, text="2) DESTINO", font=("Arial", 12, "bold")).pack(pady=(14, 5))
        self.btn_dest = ctk.CTkButton(self.sidebar, text="📍 Elegir carpeta de salida", command=self.elegir_destino, fg_color="#8e44ad", width=230)
        self.btn_dest.pack(pady=4, padx=20)
        self.lbl_destino = ctk.CTkLabel(
            self.sidebar,
            text=f"Ruta actual:\n{self.ruta_destino}",
            text_color="lightgray",
            justify="left",
            wraplength=240,
            anchor="w"
        )
        self.lbl_destino.pack(padx=20, pady=(5, 8), fill="x")
        self.btn_abrir_destino = ctk.CTkButton(self.sidebar, text="👁️ Abrir carpeta destino", command=self.abrir_destino, fg_color="#34495e", width=230)
        self.btn_abrir_destino.pack(pady=(0, 12), padx=20)

        self.idiomas = {"Español": "es", "Inglés": "en", "Francés": "fr", "Italiano": "it", "Portugués": "pt", "Alemán": "de", "Chino": "zh-CN"}
        ctk.CTkLabel(self.sidebar, text="3) TRADUCCIÓN", font=("Arial", 12, "bold")).pack(pady=(2, 5))
        self.combo_idioma = ctk.CTkComboBox(self.sidebar, values=list(self.idiomas.keys()))
        self.combo_idioma.set("Español")
        self.combo_idioma.pack(pady=(4, 12), padx=20)

        self.formato_var = ctk.StringVar(value="PDF")
        ctk.CTkRadioButton(self.sidebar, text="Exportar a PDF", variable=self.formato_var, value="PDF").pack()
        ctk.CTkRadioButton(self.sidebar, text="Exportar a Word", variable=self.formato_var, value="DOCX").pack(pady=(4, 8))

        ctk.CTkLabel(self.sidebar, text="ESTRUCTURA:", font=("Arial", 12, "bold")).pack(pady=(10, 5))
        self.estructura_var = ctk.StringVar(value="ESTRICTO")
        ctk.CTkRadioButton(
            self.sidebar,
            text="Bloque continuo",
            variable=self.estructura_var,
            value="CONTINUO"
        ).pack()
        ctk.CTkRadioButton(
            self.sidebar,
            text="Separado página por página",
            variable=self.estructura_var,
            value="ESTRICTO"
        ).pack(pady=(5, 0))

        self.btn_traducir = ctk.CTkButton(self.sidebar, text="🚀 Traducir todo", fg_color="#2ecc71", command=self.iniciar_proceso, height=48, width=230, font=("Arial", 14, "bold"))
        self.btn_traducir.pack(pady=(18, 10), padx=20)

        ctk.CTkButton(self.sidebar, text="🧹 Limpiar lista", fg_color="#e74c3c", command=self.limpiar_lista, width=230).pack()

        self.status = ctk.CTkLabel(self.sidebar, text="Listo • Doble clic para editar páginas", text_color="cyan", wraplength=240)
        self.status.pack(pady=16, padx=20)

        # 2. PANEL DERECHO (CONTENEDOR DOBLE)
        self.right_container = ctk.CTkFrame(self, fg_color="transparent")
        self.right_container.grid(row=0, column=1, sticky="nsew", padx=10, pady=10)
        self.right_container.grid_columnconfigure(0, weight=1)
        self.right_container.grid_rowconfigure(0, weight=1)

        # SUB-PANEL A: Visor de Archivos
        self.file_viewer = ctk.CTkScrollableFrame(self.right_container, label_text="DOCUMENTOS • Arrastra y suelta PDFs aquí")
        self.file_viewer.grid(row=0, column=0, sticky="nsew")
        self.file_viewer.grid_columnconfigure((0, 1, 2), weight=1)

        # SUB-PANEL B: Selector de Páginas (Oculto al inicio)
        self.page_selector = ctk.CTkFrame(self.right_container)
        # Se mostrará con .grid() cuando se necesite

        self.drop_target_register(DND_FILES)
        self.dnd_bind('<<Drop>>', self.al_soltar)

    # --- NAVEGACIÓN ENTRE PANELES ---
    def mostrar_selector(self, path):
        self.archivo_editando = path
        self.file_viewer.grid_forget() # Escondemos la lista
        self.page_selector.grid(row=0, column=0, sticky="nsew") # Mostramos el editor
        self.dibujar_selector_paginas(path)

    def volver_a_lista(self):
        self.page_selector.grid_forget()
        self.file_viewer.grid(row=0, column=0, sticky="nsew")
        self.archivo_editando = None
        self.actualizar_visor_archivos()

    # --- LÓGICA DE ARCHIVOS ---
    def agregar_a_cola(self, path):
        if path not in self.cola_archivos:
            doc = fitz.open(path)
            self.cola_archivos[path] = list(range(len(doc)))
            doc.close()

    def al_soltar(self, event):
        import re
        paths = [p[0] if p[0] else p[1] for p in re.findall(r'\{(.*?)\}|(\S+)', event.data)]
        for p in paths:
            p = p.strip('{}').replace('"', '')
            if os.path.isdir(p):
                for f in os.listdir(p):
                    if f.lower().endswith('.pdf'): self.agregar_a_cola(os.path.join(p, f))
            elif p.lower().endswith('.pdf'): self.agregar_a_cola(p)
        self.actualizar_visor_archivos()

    def actualizar_visor_archivos(self):
        for widget in self.file_viewer.winfo_children(): widget.destroy()

        if not self.cola_archivos:
            estado_vacio = ctk.CTkFrame(self.file_viewer, fg_color="#2c3e50")
            estado_vacio.grid(row=0, column=0, columnspan=3, sticky="ew", padx=24, pady=24)
            ctk.CTkLabel(estado_vacio, text="📄 No hay PDFs cargados", font=("Arial", 16, "bold")).pack(pady=(18, 8))
            ctk.CTkLabel(
                estado_vacio,
                text="Arrastra archivos PDF a esta zona\no usa los botones de la izquierda.",
                text_color="lightgray",
                justify="center",
                font=("Arial", 12)
            ).pack(pady=(0, 18))
            return

        for i, (path, seleccion) in enumerate(self.cola_archivos.items()):
            frame = ctk.CTkFrame(self.file_viewer, fg_color="#2c3e50")
            frame.grid(row=i//3, column=i%3, padx=10, pady=10, sticky="nsew")
            
            doc = fitz.open(path)
            pix = doc[0].get_pixmap(matrix=fitz.Matrix(0.1, 0.1))
            img = ctk.CTkImage(Image.frombytes("RGB", [pix.width, pix.height], pix.samples), size=(100, 130))
            
            lbl_img = ctk.CTkLabel(frame, image=img, text="")
            lbl_img.pack(pady=5)
            ctk.CTkLabel(
                frame,
                text=os.path.basename(path),
                font=("Arial", 10, "bold"),
                wraplength=140,
                justify="center"
            ).pack(padx=6)
            ctk.CTkLabel(frame, text=f"Páginas seleccionadas: {len(seleccion)}/{len(doc)}", font=("Arial", 9), text_color="gray").pack(pady=(2, 2))
            ctk.CTkLabel(frame, text="Doble clic para editar", font=("Arial", 9), text_color="lightgray").pack(pady=(0, 6))
            doc.close()

            # Doble clic para cambiar de panel
            frame.bind("<Double-Button-1>", lambda e, p=path: self.mostrar_selector(p))
            lbl_img.bind("<Double-Button-1>", lambda e, p=path: self.mostrar_selector(p))

    # --- DIBUJAR EDITOR DE PÁGINAS ---
    def dibujar_selector_paginas(self, path):
        for widget in self.page_selector.winfo_children(): widget.destroy()
        
        lbl_info = ctk.CTkLabel(self.page_selector, text=f"Selecciona páginas • {os.path.basename(path)}", font=("Arial", 14, "bold"))
        lbl_info.pack(pady=(10, 4))
        ctk.CTkLabel(self.page_selector, text="Marca las páginas que quieras traducir y guarda los cambios.", text_color="lightgray").pack(pady=(0, 8))

        scroll_p = ctk.CTkScrollableFrame(self.page_selector)
        scroll_p.pack(fill="both", expand=True, padx=10, pady=5)
        scroll_p.grid_columnconfigure((0,1,2,3), weight=1)

        doc = fitz.open(path)
        checks = {}
        for i in range(len(doc)):
            f_p = ctk.CTkFrame(scroll_p, fg_color="#333333")
            f_p.grid(row=i//4, column=i%4, padx=5, pady=5)
            
            pix = doc[i].get_pixmap(matrix=fitz.Matrix(0.08, 0.08))
            img = ctk.CTkImage(Image.frombytes("RGB", [pix.width, pix.height], pix.samples), size=(80, 110))
            
            ctk.CTkLabel(f_p, image=img, text="").pack()
            var = ctk.BooleanVar(value=i in self.cola_archivos[path])
            checks[i] = var
            ctk.CTkCheckBox(f_p, text=f"Pág {i+1}", variable=var).pack(pady=2)

        def confirmar():
            self.cola_archivos[path] = sorted({pg for pg, v in checks.items() if v.get()})
            doc.close()
            self.volver_a_lista()

        btn_confirmar = ctk.CTkButton(self.page_selector, text="✅ Guardar selección y volver", fg_color="#2ecc71", command=confirmar, height=42)
        btn_confirmar.pack(pady=10)

    # --- PROCESO TRADUCCIÓN ---
    def iniciar_proceso(self):
        if not self.cola_archivos:
            return
        config = {
            "idioma_destino": self.idiomas[self.combo_idioma.get()],
            "formato": self.formato_var.get(),
            "separar_paginas": self.estructura_var.get() == "ESTRICTO"
        }
        self._mostrar_ventana_progreso_ui()
        self.btn_traducir.configure(state="disabled")
        threading.Thread(target=self.hilo_traduccion, args=(config,), daemon=True).start()

    def _set_status_ui(self, texto, color="gray"):
        self.after(0, lambda: self.status.configure(text=texto, text_color=color))

    def _mostrar_ventana_progreso_ui(self):
        if self.progress_window and self.progress_window.winfo_exists():
            if self.progress_bar_popup:
                self.progress_bar_popup.set(0)
            if self.progress_label_popup:
                self.progress_label_popup.configure(text="Traduciendo... 0%")
            return

        self.attributes("-alpha", 1.0)

        win = ctk.CTkToplevel(self)
        win.title("Traduciendo")
        win.geometry("420x150")
        win.resizable(False, False)
        win.protocol("WM_DELETE_WINDOW", lambda: None)
        win.lift()
        win.attributes("-topmost", True)
        self.after(300, lambda: win.attributes("-topmost", False) if win.winfo_exists() else None)

        self.update_idletasks()
        x = self.winfo_rootx() + (self.winfo_width() // 2) - 210
        y = self.winfo_rooty() + (self.winfo_height() // 2) - 75
        win.geometry(f"420x150+{x}+{y}")

        ctk.CTkLabel(win, text="Traduciendo documento(s)", font=("Arial", 16, "bold")).pack(pady=(18, 6))
        self.progress_label_popup = ctk.CTkLabel(win, text="Traduciendo... 0%", text_color="lightgray")
        self.progress_label_popup.pack(pady=(0, 8))
        self.progress_bar_popup = ctk.CTkProgressBar(win, width=340)
        self.progress_bar_popup.pack(pady=(0, 14))
        self.progress_bar_popup.set(0)

        win.grab_set()
        self.progress_window = win

    def _cerrar_ventana_progreso_ui(self):
        if self.progress_window and self.progress_window.winfo_exists():
            try:
                self.progress_window.grab_release()
            except Exception:
                pass
            self.progress_window.destroy()
        self.progress_window = None
        self.progress_bar_popup = None
        self.progress_label_popup = None
        self.attributes("-alpha", 1.0)

    def _actualizar_progreso_popup_ui(self, valor):
        if self.progress_bar_popup and self.progress_bar_popup.winfo_exists():
            self.progress_bar_popup.set(valor)
        if self.progress_label_popup and self.progress_label_popup.winfo_exists():
            self.progress_label_popup.configure(text=f"Traduciendo... {int(valor * 100)}%")

    def _set_progress_ui(self, valor):
        v = max(0.0, min(1.0, valor))
        self.after(0, lambda: self._actualizar_progreso_popup_ui(v))

    def hilo_traduccion(self, config):
        try:
            self._set_progress_ui(0)
            translator = GoogleTranslator(source='auto', target=config["idioma_destino"])

            archivos = list(self.cola_archivos.items())
            total_bloques = 0
            estructura_archivos = []

            for p_in, paginas in archivos:
                if not paginas:
                    continue
                reader = PdfReader(p_in)
                doc_fitz = fitz.open(p_in)
                paginas_origen = []
                paginas_ordenadas = sorted(set(paginas))

                for idx in paginas_ordenadas:
                    parrafos = self.extraer_parrafos_pagina(doc_fitz, idx)
                    if not parrafos:
                        t = reader.pages[idx].extract_text()
                        parrafos = self.extraer_parrafos(t) if t else []

                    paginas_origen.append(parrafos)
                    for p in parrafos:
                        total_bloques += max(1, len(self.construir_bloques_traduccion(p)))

                estructura_archivos.append((p_in, paginas_origen))
                doc_fitz.close()

            if total_bloques <= 0:
                total_bloques = 1

            bloques_completados = 0

            for i, (p_in, paginas_origen) in enumerate(estructura_archivos):
                self._set_status_ui(f"Traduciendo {i+1}/{len(estructura_archivos)}...", "orange")
                paginas_traducidas = []
                separar_paginas = config["separar_paginas"]

                def avanzar_progreso():
                    nonlocal bloques_completados
                    bloques_completados += 1
                    self._set_progress_ui(bloques_completados / total_bloques)

                for parrafos in paginas_origen:
                    parrafos_traducidos = []
                    for p in parrafos:
                        parrafos_traducidos.append(
                            self.traducir_texto(translator, p, progreso_callback=avanzar_progreso)
                        )
                    paginas_traducidas.append(parrafos_traducidos)

                ext = ".pdf" if config["formato"] == "PDF" else ".docx"
                p_out = os.path.join(self.ruta_destino, f"Traducido_{os.path.basename(p_in).replace('.pdf', ext)}")
                
                if config["formato"] == "PDF":
                    self.guardar_pdf_estructurado(paginas_traducidas, p_out, separar_paginas=separar_paginas)
                else:
                    self.guardar_docx(paginas_traducidas, p_out, separar_paginas=separar_paginas)

            self._set_progress_ui(1)
            self._set_status_ui("¡COMPLETADO!", "green")
            self.after(0, lambda: messagebox.showinfo("Éxito", "Archivos guardados correctamente."))
        except Exception as e:
            self._set_status_ui("Error en la traducción", "red")
            self.after(0, lambda: messagebox.showerror("Error", f"No se pudo completar la traducción:\n{e}"))
        finally:
            self.after(0, self._cerrar_ventana_progreso_ui)
            self.after(0, lambda: self.btn_traducir.configure(state="normal"))

    def guardar_pdf(self, texto, path):
        c = canvas.Canvas(path, pagesize=A4)
        text_obj = c.beginText(50, 800); text_obj.setFont("Helvetica", 10)
        for linea in texto.split('\n'):
            if text_obj.getY() < 50:
                c.drawText(text_obj); c.showPage()
                text_obj = c.beginText(50, 800); text_obj.setFont("Helvetica", 10)
            text_obj.textLine(linea[:95])
        c.drawText(text_obj); c.save()

    def guardar_pdf_estructurado(self, paginas_texto, path, separar_paginas=True):
        c = canvas.Canvas(path, pagesize=A4)
        text_obj = c.beginText(50, 800)
        text_obj.setFont("Helvetica", 10)
        pagina_vacia = True

        for i, parrafos_pagina in enumerate(paginas_texto):
            if parrafos_pagina:
                for parrafo in parrafos_pagina:
                    texto = parrafo.strip()
                    if not texto:
                        continue

                    lineas = textwrap.wrap(texto, width=95, break_long_words=False, break_on_hyphens=False)
                    if not lineas:
                        lineas = [""]

                    for linea in lineas:
                        if text_obj.getY() < 50:
                            c.drawText(text_obj)
                            c.showPage()
                            text_obj = c.beginText(50, 800)
                            text_obj.setFont("Helvetica", 10)
                            pagina_vacia = True
                        text_obj.textLine(linea)
                        pagina_vacia = False

                    if text_obj.getY() < 50:
                        c.drawText(text_obj)
                        c.showPage()
                        text_obj = c.beginText(50, 800)
                        text_obj.setFont("Helvetica", 10)
                        pagina_vacia = True
                    text_obj.textLine("")
                    pagina_vacia = False

            if separar_paginas and i < len(paginas_texto) - 1:
                if not pagina_vacia:
                    c.drawText(text_obj)
                    c.showPage()
                    text_obj = c.beginText(50, 800)
                    text_obj.setFont("Helvetica", 10)
                    pagina_vacia = True

        c.drawText(text_obj)
        c.save()

    def construir_bloques_traduccion(self, texto, tam_bloque=1800):
        contenido = texto.strip()
        if contenido.startswith(("- ", "• ", "* ")):
            contenido = contenido[2:].strip()

        if not contenido:
            return []

        oraciones = self.dividir_en_oraciones(contenido)
        if not oraciones:
            oraciones = [contenido]

        bloques = []
        actual = ""
        for oracion in oraciones:
            candidato = (actual + " " + oracion).strip() if actual else oracion
            if len(candidato) <= tam_bloque:
                actual = candidato
            else:
                if actual:
                    bloques.append(actual)
                actual = oracion

        if actual:
            bloques.append(actual)

        return bloques

    def traducir_texto(self, translator, texto, tam_bloque=1800, progreso_callback=None):
        if not texto.strip():
            return ""

        prefijo = ""
        contenido = texto.strip()
        if contenido.startswith(("- ", "• ", "* ")):
            prefijo = contenido[:2]
            contenido = contenido[2:].strip()

        if not contenido:
            return prefijo.strip()

        bloques = self.construir_bloques_traduccion(contenido, tam_bloque=tam_bloque)

        traducido = []
        for bloque in bloques:
            try:
                traducido.append(translator.translate(bloque))
            except Exception:
                traducido.append(bloque)
            if progreso_callback:
                progreso_callback()

        texto_traducido = " ".join(traducido).strip()
        texto_traducido = self.limpiar_texto_traducido(texto_traducido)
        return f"{prefijo}{texto_traducido}" if prefijo else texto_traducido

    def dividir_en_oraciones(self, texto):
        texto_norm = re.sub(r"\s+", " ", texto).strip()
        if not texto_norm:
            return []

        partes = re.split(r'(?<=[.!?;:])\s+(?=[A-ZÁÉÍÓÚÜÑ0-9"“(])', texto_norm)
        if len(partes) == 1:
            return [texto_norm]
        return [p.strip() for p in partes if p.strip()]

    def limpiar_texto_traducido(self, texto):
        t = texto.strip()
        t = re.sub(r"\s+", " ", t)
        t = re.sub(r"\s+([,.;:!?])", r"\1", t)
        t = re.sub(r"([¿¡])\s+", r"\1", t)
        t = re.sub(r"\(\s+", "(", t)
        t = re.sub(r"\s+\)", ")", t)
        return t

    def extraer_parrafos_pagina(self, doc_fitz, page_index):
        page = doc_fitz[page_index]
        data = page.get_text("dict")
        parrafos = []

        blocks = data.get("blocks", [])
        blocks_ordenados = sorted(blocks, key=lambda b: (round(b.get("bbox", [0, 0])[1], 1), round(b.get("bbox", [0, 0])[0], 1)))

        for bloque in blocks_ordenados:
            if bloque.get("type") != 0:
                continue

            lineas = []
            for linea in bloque.get("lines", []):
                spans = [s.get("text", "").strip() for s in linea.get("spans", [])]
                texto_linea = " ".join([s for s in spans if s]).strip()
                if not texto_linea:
                    continue

                bbox = linea.get("bbox", [0, 0, 0, 0])
                lineas.append({
                    "texto": texto_linea,
                    "y0": bbox[1],
                    "y1": bbox[3]
                })

            if not lineas:
                continue

            actual = ""
            prev = None
            for linea in lineas:
                texto_linea = linea["texto"]
                if not actual:
                    actual = texto_linea
                    prev = linea
                    continue

                altura_prev = max((prev["y1"] - prev["y0"]), 1)
                gap = linea["y0"] - prev["y1"]
                nueva_seccion = gap > (altura_prev * 0.75)

                es_lista = self.es_linea_lista(texto_linea)
                if nueva_seccion or es_lista:
                    parrafos.append(actual.strip())
                    actual = texto_linea
                else:
                    if actual.endswith("-"):
                        actual = actual[:-1] + texto_linea
                    else:
                        actual += " " + texto_linea

                prev = linea

            if actual.strip():
                parrafos.append(actual.strip())

        return [p for p in parrafos if p.strip()]

    def es_linea_lista(self, texto):
        t = texto.strip()
        if t.startswith(("- ", "• ", "* ")):
            return True
        return len(t) > 2 and t[0].isdigit() and t[1:3] in (". ", ") ")

    def extraer_parrafos(self, texto_pagina):
        lineas = [l.strip() for l in texto_pagina.splitlines()]
        parrafos = []
        actual = ""

        for linea in lineas:
            if not linea:
                if actual:
                    parrafos.append(actual.strip())
                    actual = ""
                continue

            es_lista = linea.startswith(("- ", "• "))
            es_lista_numerada = len(linea) > 2 and linea[0].isdigit() and linea[1:3] in (". ", ") ")
            if es_lista or es_lista_numerada:
                if actual:
                    parrafos.append(actual.strip())
                    actual = ""
                parrafos.append(linea)
                continue

            if not actual:
                actual = linea
                continue

            if actual.endswith("-"):
                actual = actual[:-1] + linea
            else:
                actual += " " + linea

        if actual:
            parrafos.append(actual.strip())

        return parrafos

    def guardar_docx(self, paginas_texto, path, separar_paginas=True):
        doc = Document()
        for i, parrafos_pagina in enumerate(paginas_texto):
            if separar_paginas and i > 0:
                doc.add_page_break()

            if not parrafos_pagina:
                doc.add_paragraph("")
                continue

            for parrafo in parrafos_pagina:
                texto = parrafo.strip()
                if not texto:
                    continue

                if texto.startswith(("- ", "• ")):
                    doc.add_paragraph(texto[2:].strip(), style="List Bullet")
                elif len(texto) > 2 and texto[0].isdigit() and texto[1:3] in (". ", ") "):
                    contenido = texto[3:].strip() if len(texto) > 3 else texto
                    doc.add_paragraph(contenido, style="List Number")
                else:
                    doc.add_paragraph(texto)

        doc.save(path)

    def examinar_archivos(self):
        files = filedialog.askopenfilenames(filetypes=[("PDF", "*.pdf")])
        if files:
            for f in files: self.agregar_a_cola(f)
            self.actualizar_visor_archivos()

    def examinar_carpeta(self):
        folder = filedialog.askdirectory()
        if folder:
            for f in os.listdir(folder):
                if f.lower().endswith('.pdf'): self.agregar_a_cola(os.path.join(folder, f))
            self.actualizar_visor_archivos()

    def elegir_destino(self):
        dest = filedialog.askdirectory()
        if dest:
            self.ruta_destino = dest
            self.actualizar_label_destino()

    def actualizar_label_destino(self):
        self.lbl_destino.configure(text=f"Ruta actual:\n{self.ruta_destino}")

    def abrir_destino(self):
        if not os.path.isdir(self.ruta_destino):
            messagebox.showerror("Destino no válido", "La carpeta de destino no existe.")
            return
        try:
            os.startfile(self.ruta_destino)
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo abrir la carpeta destino:\n{e}")

    def limpiar_lista(self):
        self.cola_archivos = {}; self.actualizar_visor_archivos()

if __name__ == "__main__":
    app = App()
    app.mainloop()
